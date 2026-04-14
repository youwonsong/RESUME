from docx import Document

def normalize(text):
    return text.strip().lower()

def is_list_paragraph(para):
    """Return True if paragraph is part of a numbered/bulleted list."""
    return (
        para._p.pPr is not None and
        para._p.pPr.numPr is not None
    )

def get_heading_level(style_name):
    """Extract heading level if it exists."""
    if "Heading" in style_name:
        parts = style_name.split()
        for p in parts:
            if p.isdigit():
                return int(p)
    return None


def extract_section(docx_path, section_title, output_path):

    doc = Document(docx_path)
    new_doc = Document()

    target = normalize(section_title)

    inside_section = False
    target_level = None

    for para in doc.paragraphs:

        text = normalize(para.text)
        style = para.style.name

        is_heading = "Heading" in style
        is_list = is_list_paragraph(para)

        # treat either heading or list paragraph as possible section header
        if is_heading or is_list:

            level = get_heading_level(style) if is_heading else 99

            # stop when next heading of same or higher level appears
            if inside_section and level <= target_level:
                break

            # match section title
            if target in text:
                inside_section = True
                target_level = level
                continue

        if inside_section:
            new_doc.add_paragraph(para.text)

    if inside_section:
        new_doc.save(str(output_path))
        print("Section extracted successfully.")
    else:
        print("Section not found.")