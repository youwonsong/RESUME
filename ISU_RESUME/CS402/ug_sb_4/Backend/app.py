from flask import Flask, send_from_directory, request, jsonify
import os
from add_data import build_context, Extractor
from docxtpl import DocxTemplate
from docx import Document
from extract_section import extract_section
import time
import traceback
import subprocess
from validation import validation

app = Flask(
    __name__,
    static_folder="../Frontend",
    static_url_path=""
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "spreadsheets")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
DOCX_UPLOAD_DIR = os.path.join(BASE_DIR, "uploaded_docs")
PDF_UPLOAD_DIR = os.path.join(BASE_DIR, "uploaded_pdfs")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(DOCX_UPLOAD_DIR, exist_ok=True)
os.makedirs(PDF_UPLOAD_DIR, exist_ok=True)

@app.route("/")
def index():
    return app.send_static_file("index.html")

@app.route("/generate-page")
def generate_page():
    return app.send_static_file("generate.html")

@app.route("/extract-page")
def extract_page():
    return app.send_static_file("extract.html")

@app.route("/validate-page")
def validate_page():
    return app.send_static_file("validate.html")

@app.route("/upload", methods=["POST"])
def upload():
    file = request.files["file"]
    filepath = os.path.join(UPLOAD_DIR, file.filename)
    file.save(filepath)
    return jsonify({"message": "File uploaded.", "filename": file.filename})

@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json()
    filename = data.get("filename")
    template_name = data.get("template")

    if not filename:
        return jsonify({"error": "No filename provided."}), 400
    if not template_name:
        return jsonify({"error": "No template selected."}), 400

    spreadsheet_path = os.path.join(UPLOAD_DIR, filename)

    template_path = os.path.join(BASE_DIR, "templates", template_name)

    spreadsheet_name = os.path.splitext(filename)[0]
    template_base = os.path.splitext(template_name)[0]

    output_filename = f"{spreadsheet_name}_{template_base}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    try:
        extract = Extractor(spreadsheet_path)

        def normalize(method):
            def wrapper():
                result = method()

                if isinstance(result, tuple):
                    if len(result) == 2 and hasattr(result[1], "iloc"):
                        return result

                    years = None
                    dataframe = None

                    for item in result:
                        if isinstance(item, int) and years is None:
                            years = item
                        elif hasattr(item, "iloc") and dataframe is None:
                            dataframe = item

                    if dataframe is not None:
                        return (years, dataframe)

                if hasattr(result, "iloc"):
                    return (None, result)

                return result
            return wrapper

        extract.grab_key_personnel = normalize(extract.grab_key_personnel)
        extract.grab_other_personnel = normalize(extract.grab_other_personnel)
        extract.grab_benefits = normalize(extract.grab_benefits)
        extract.grab_direct_cost = normalize(extract.grab_direct_cost)
        extract.grab_domestic_travel = normalize(extract.grab_domestic_travel)
        
        context = build_context(extract)

        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)

        pdf_filename = output_filename.replace(".docx", ".pdf")

        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            output_path,
            "--outdir", OUTPUT_DIR
        ], check=True)

        return jsonify({"message": "Document generated.", "generatedName": output_filename, "pdfName": pdf_filename})
    except Exception as e:
        print("ERROR in /generate:", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/extract-section", methods=["POST"])
def extract_section_route():
    file = request.files.get("file")
    section_title = request.form.get("sectionTitle")

    if not file:
        return jsonify({"error": "No document uploaded."}), 400
    if not section_title:
        return jsonify({"error": "No section title provided."}), 400

    input_filename = f"source_{int(time.time())}_{file.filename}"
    input_path = os.path.join(DOCX_UPLOAD_DIR, input_filename)
    file.save(input_path)

    safe_section = section_title.replace(" ", "_")
    original_name = os.path.splitext(file.filename)[0]

    output_filename = f"{original_name}_{safe_section}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    try:
        extract_section(input_path, section_title, output_path)

        if not os.path.exists(output_path):
            return jsonify({"error": "Section not found in document."}), 400

        pdf_filename = None
        pdf_path = os.path.join(OUTPUT_DIR, output_filename.replace(".docx", ".pdf"))

        try:
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                output_path,
                "--outdir", OUTPUT_DIR
            ], check=True)

            if os.path.exists(pdf_path):
                pdf_filename = os.path.basename(pdf_path)

        except Exception as e:
            print("PDF conversion failed for extracted section:", e)

        return jsonify({
            "message": "Section extracted successfully.",
            "generatedName": output_filename,
            "pdfName": pdf_filename
        })
    except Exception as e:
        print("ERROR in /extract-section:", e)
        return jsonify({"error": str(e)}), 500

@app.route("/validate", methods=["POST"])
def validate_documents():
    spreadsheet = request.files.get("spreadsheet")
    streamlyne_pdf = request.files.get("streamlynePdf")

    if not spreadsheet:
        return jsonify({"error": "No spreadsheet uploaded."}), 400

    if not streamlyne_pdf:
        return jsonify({"error": "No Streamlyne PDF uploaded."}), 400

    spreadsheet_path = os.path.join(UPLOAD_DIR, spreadsheet.filename)
    pdf_path = os.path.join(PDF_UPLOAD_DIR, streamlyne_pdf.filename)

    spreadsheet.save(spreadsheet_path)
    streamlyne_pdf.save(pdf_path)

    spreadsheet_base = os.path.splitext(spreadsheet.filename)[0]
    output_filename = f"{spreadsheet_base}_validation_report.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    pdf_filename = output_filename.replace(".docx", ".pdf")

    try:
        salaries_df, benefits_df, direct_cost_df, travel_df, indirect_df = validation(
            spreadsheet_path,
            pdf_path
        )

        create_validation_report(
            output_path,
            spreadsheet.filename,
            streamlyne_pdf.filename,
            salaries_df,
            benefits_df,
            direct_cost_df,
            travel_df,
            indirect_df
        )

        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            output_path,
            "--outdir", OUTPUT_DIR
        ], check=True)

        return jsonify({
            "message": "Validation report generated.",
            "generatedName": output_filename,
            "pdfName": pdf_filename
        })
    except Exception as e:
        print("ERROR in /validate:", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

def add_dataframe_table(doc, title, df):
    doc.add_heading(title, level=2)

    if df.empty:
        doc.add_paragraph("No results found.")
        return

    columns_to_keep = ["spreadsheet_name"]

    for col in df.columns:
        if col.startswith("Difference Year"):
            columns_to_keep.append(col)

    filtered_df = df[columns_to_keep].copy()

    table = doc.add_table(rows=1, cols=len(filtered_df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, col_name in enumerate(filtered_df.columns):
        header_cells[i].text = str(col_name)

    for _, row in filtered_df.iterrows():
        row_cells = table.add_row().cells

        for i, col_name in enumerate(filtered_df.columns):
            value = row[col_name]
            text = str(value)

            if col_name.startswith("Difference Year"):
                try:
                    numeric_value = float(value)
                    text = str(round(numeric_value, 2))
                except:
                    numeric_value = None

                paragraph = row_cells[i].paragraphs[0]
                run = paragraph.add_run(text)

                if numeric_value is not None and numeric_value != 0:
                    run.bold = True
            else:
                row_cells[i].text = text

def create_validation_report(output_path, spreadsheet_name, pdf_name, salaries_df, benefits_df, direct_cost_df, travel_df, indirect_df):
    doc = Document()

    doc.add_heading("Validation Report", 0)
    doc.add_paragraph(f"Spreadsheet: {spreadsheet_name}")
    doc.add_paragraph(f"Streamlyne PDF: {pdf_name}")

    add_dataframe_table(doc, "Salaries", salaries_df)
    add_dataframe_table(doc, "Benefits", benefits_df)
    add_dataframe_table(doc, "Direct Costs", direct_cost_df)
    add_dataframe_table(doc, "Travel", travel_df)
    add_dataframe_table(doc, "Indirect Costs", indirect_df)

    doc.save(output_path)

@app.route("/documents", methods=["GET"])
def list_documents():
    documents = []

    for filename in os.listdir(OUTPUT_DIR):
        if filename.lower().endswith(".docx"):
            filepath = os.path.join(OUTPUT_DIR, filename)

            pdf_name = filename.replace(".docx", ".pdf")
            pdf_path = os.path.join(OUTPUT_DIR, pdf_name)

            documents.append({
                "name": filename,
                "generatedName": filename,
                "pdfName": pdf_name if os.path.exists(pdf_path) else None,
                "date": os.path.getmtime(filepath)
            })

    documents.sort(key=lambda doc: doc["date"], reverse=True)
    return jsonify(documents)

@app.route("/download/<filename>")
def download(filename):
    return send_from_directory(OUTPUT_DIR, filename)

@app.route("/delete/<filename>", methods=["DELETE"])
def delete_document(filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    pdf_path = os.path.join(OUTPUT_DIR, filename.replace(".docx", ".pdf"))

    if not os.path.exists(filepath):
        return jsonify({"error": "File not found."}), 404

    try:
        os.remove(filepath)

        if os.path.exists(pdf_path):
            os.remove(pdf_path)

        return jsonify({"message": "File deleted."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080, debug=True)